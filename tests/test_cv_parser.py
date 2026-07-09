"""CV ayristirma servisi testleri (Kisi 2 - Hafta 2).

Gemini API cagirmaz; tamamen offline calisir.
Test dosyalari diske checked-in binary olarak konmaz, test icinde uretilir.
"""

import io

import pytest
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from services.cv_parser import CVParseError, extract_text


def _pdf_bytes(*lines: str) -> bytes:
    """Verilen satirlari iceren gercek bir PDF uretir."""
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    y = 800
    for line in lines:
        pdf.drawString(72, y, line)
        y -= 20
    pdf.save()
    return buffer.getvalue()


def _docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """Verilen paragraf ve tablo hucrelerini iceren gercek bir DOCX uretir."""
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for cell_index, cell_text in enumerate(row):
                table.cell(row_index, cell_index).text = cell_text
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- Mutlu yol: uc formatta da metin cikar ---


def test_pdf_metni_cikarilir():
    data = _pdf_bytes("Ahmet Yilmaz", "Backend Developer", "Python ve FastAPI deneyimi")

    result = extract_text(data, "cv.pdf")

    assert "Ahmet Yilmaz" in result
    assert "Python ve FastAPI deneyimi" in result


def test_docx_paragraf_metni_cikarilir():
    data = _docx_bytes(["Ayse Demir", "5 yil Frontend deneyimi"])

    result = extract_text(data, "cv.docx")

    assert "Ayse Demir" in result
    assert "5 yil Frontend deneyimi" in result


def test_docx_tablo_hucreleri_de_cikarilir():
    """CV'lerin cogu Word'de tabloyla dizilir; doc.paragraphs bunlari atlar."""
    data = _docx_bytes(
        paragraphs=["Mehmet Kaya"],
        table_rows=[["Beceriler", "Docker, Kubernetes"], ["Diller", "Python, Go"]],
    )

    result = extract_text(data, "cv.docx")

    assert "Mehmet Kaya" in result
    assert "Docker, Kubernetes" in result
    assert "Python, Go" in result


def test_txt_metni_cikarilir():
    data = "Zeynep Sahin\nData Scientist\nPandas, scikit-learn".encode("utf-8")

    result = extract_text(data, "cv.txt")

    assert "Zeynep Sahin" in result
    assert "Pandas, scikit-learn" in result


def test_uzanti_buyuk_harf_olsa_da_calisir():
    data = "Bu bir CV metnidir ve yeterince uzundur.".encode("utf-8")

    assert "CV metnidir" in extract_text(data, "CV.TXT")


# --- Kodlama ---


def test_latin5_kodlu_txt_okunabilir():
    """Turkce karakterli cp1254 baytlari utf-8 olarak cozulemez; fallback devreye girmeli."""
    data = "Görkem Çetin - Yazılım Mühendisi, İstanbul".encode("cp1254")
    with pytest.raises(UnicodeDecodeError):
        data.decode("utf-8")  # testin gercekten fallback yolunu zorladiginin kaniti

    result = extract_text(data, "cv.txt")

    assert "Yazılım Mühendisi" in result


# --- Normalize ---


def test_fazla_bos_satirlar_teke_iner():
    data = "Ad Soyad\n\n\n\n\nBackend Developer pozisyonu icin basvuru".encode("utf-8")

    result = extract_text(data, "cv.txt")

    assert "\n\n\n" not in result
    assert "Ad Soyad\n\nBackend Developer" in result


def test_satir_sonu_bosluklari_kirpilir():
    data = "Ad Soyad   \r\nBackend Developer pozisyonu icin basvuru   ".encode("utf-8")

    result = extract_text(data, "cv.txt")

    assert "\r" not in result
    assert "Ad Soyad\nBackend Developer" in result


# --- Hata yolu: hepsi CVParseError ---


def test_desteklenmeyen_uzanti_hata_verir():
    with pytest.raises(CVParseError, match="Desteklenmeyen dosya turu"):
        extract_text(b"herhangi bir icerik", "cv.png")


def test_uzantisiz_dosya_hata_verir():
    with pytest.raises(CVParseError, match="Desteklenmeyen dosya turu"):
        extract_text(b"herhangi bir icerik", "cv")


def test_bos_dosya_hata_verir():
    with pytest.raises(CVParseError, match="metin cikarilamadi"):
        extract_text(b"", "cv.txt")


def test_metni_olmayan_pdf_hata_verir():
    """Taranmis PDF: sayfalari var ama metin katmani yok."""
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    pdf.showPage()  # bos sayfa, hic metin cizilmedi
    pdf.save()

    with pytest.raises(CVParseError, match="taranmis"):
        extract_text(buffer.getvalue(), "taranmis_cv.pdf")


def test_cok_kisa_metin_hata_verir():
    with pytest.raises(CVParseError, match="metin cikarilamadi"):
        extract_text("Merhaba".encode("utf-8"), "cv.txt")


def test_bozuk_pdf_ham_kutuphane_hatasi_sizdirmaz():
    with pytest.raises(CVParseError, match="PDF okunamadi"):
        extract_text(b"bu bir PDF degil, sadece rastgele baytlar", "cv.pdf")


def test_bozuk_docx_ham_kutuphane_hatasi_sizdirmaz():
    with pytest.raises(CVParseError, match="DOCX okunamadi"):
        extract_text(b"bu bir DOCX degil, sadece rastgele baytlar", "cv.docx")
