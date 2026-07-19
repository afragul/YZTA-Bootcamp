"""CV dosyalarindan duz metin cikaran servis (Kisi 2 - Hafta 2).

Tek giris noktasi: extract_text(data, filename).
Cagiran taraf dosyayi nasil aldigini (UploadFile, disk, indirme) onemsemez;
elinde bytes ve dosya adi olmasi yeterlidir.
"""

import io
import re
from pathlib import Path

import pdfplumber
from docx import Document

# Bu esigin altindaki metin "CV degil, bos" sayilir. Bilincli olarak dusuk:
# "bu bir CV mi?" karari parser'in degil, LLM analizinin isi.
MIN_TEXT_LENGTH = 30

# Turkce Windows'ta olusturulmus dosyalarda utf-8 disi kodlama sik goruluyor.
TEXT_ENCODINGS = ("utf-8", "cp1254")


class CVParseError(Exception):
    """CV dosyasindan kullanilabilir metin cikarilamadigini belirtir."""


def extract_text(data: bytes, filename: str) -> str:
    """Dosya icerigini duz metne cevirir.

    Args:
        data: Dosyanin ham baytlari.
        filename: Uzantisi icin kullanilan dosya adi (icerigi okunmaz).

    Returns:
        Normalize edilmis duz metin.

    Raises:
        CVParseError: Uzanti desteklenmiyorsa, dosya bozuksa veya
            anlamli metin cikarilamadiysa.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        raw_text = _parse_pdf(data)
    elif suffix == ".docx":
        raw_text = _parse_docx(data)
    elif suffix == ".txt":
        raw_text = _parse_txt(data)
    else:
        raise CVParseError(f"Desteklenmeyen dosya turu: {suffix or '(uzanti yok)'}")

    text = _normalize(raw_text)

    if len(text) < MIN_TEXT_LENGTH:
        if suffix == ".pdf":
            raise CVParseError(
                "PDF'ten metin cikarilamadi (taranmis goruntu olabilir). "
                "Lutfen metin katmani iceren bir PDF yukleyin."
            )
        raise CVParseError("Dosyadan metin cikarilamadi: icerik bos veya cok kisa.")

    return text


def _parse_pdf(data: bytes) -> str:
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
    except Exception as exc:
        raise CVParseError(f"PDF okunamadi: {exc}") from exc
    return "\n".join(pages)


def _parse_docx(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise CVParseError(f"DOCX okunamadi: {exc}") from exc

    parts = [paragraph.text for paragraph in document.paragraphs]

    # CV'lerin cogu Word'de tabloyla dizilir; doc.paragraphs tablolari atlar.
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)

    return "\n".join(parts)


def _parse_txt(data: bytes) -> str:
    for encoding in TEXT_ENCODINGS:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    # Hicbiri tutmadiysa metni kaybetmektense bozuk karakterle devam et.
    return data.decode("utf-8", errors="replace")


def _normalize(text: str) -> str:
    """Satir sonlarini ve fazla bosluklari duzenler.

    PDF'ten gelen sutun karisikligini duzeltmeye calismaz; agresif temizlik
    bilgi kaybettirir ve LLM dagink metni zaten tolere eder.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(line.strip() for line in text.split("\n"))
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
